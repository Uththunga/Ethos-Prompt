"""
Document Extractors - Extract text content from various file formats
"""
import logging
import io
import re
from typing import Dict, Any, Optional, List, Union
from dataclasses import dataclass
from datetime import datetime
import mimetypes

# Document processing imports (conditional)
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown  # type: ignore
    from bs4 import BeautifulSoup
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False


# Optional advanced PDF processing libs (conditional)
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

logger = logging.getLogger(__name__)

@dataclass
class ExtractedDocument:
    content: str
    metadata: Dict[str, Any]
    file_type: str
    file_size: int
    extraction_time: float
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None

@dataclass
class ExtractionResult:
    success: bool
    document: Optional[ExtractedDocument] = None
    error: Optional[str] = None
    warnings: Optional[List[str]] = None

class DocumentExtractor:
    """
    Base class for document extractors
    """

    def __init__(self):
        self.supported_types = []
        self.max_file_size = 50 * 1024 * 1024  # 50MB default

    def can_extract(self, file_type: str, mime_type: Optional[str] = None) -> bool:
        """Check if this extractor can handle the file type"""
        return file_type.lower() in self.supported_types

    def extract(self, file_content: bytes, filename: Optional[str] = None) -> ExtractionResult:
        """Extract content from file"""
        raise NotImplementedError("Subclasses must implement extract method")

    def _detect_encoding(self, content: bytes) -> str:
        """Detect text encoding"""
        if CHARDET_AVAILABLE:
            try:
                result = chardet.detect(content)
                return str(result.get('encoding') or 'utf-8')
            except Exception:
                pass
        return 'utf-8'

    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        return text.strip()

    def _calculate_stats(self, content: str) -> Dict[str, int]:
        """Calculate text statistics"""
        return {
            'character_count': len(content),
            'word_count': len(content.split()),
            'line_count': len(content.splitlines())
        }

class PDFExtractor(DocumentExtractor):
    """
    Extract text from PDF files with optional advanced handling for
    - Encrypted PDFs (best-effort empty-password decrypt)
    - Page-level extraction with per-page fallbacks
    - Optional OCR/table extraction if pdfplumber+pytesseract are available
    """

    def __init__(self):
        super().__init__()
        self.supported_types = ['pdf']

    def _ocr_page_with_plumber(self, pdf_bytes: bytes, page_index: int) -> str:
        """Try OCR on a single page using pdfplumber + pytesseract when available."""
        if not (PDFPLUMBER_AVAILABLE and PYTESSERACT_AVAILABLE):
            return ''
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                if page_index < 0 or page_index >= len(pdf.pages):
                    return ''
                page = pdf.pages[page_index]
                # Render to image and run OCR
                img = page.to_image(resolution=200).original  # PIL Image
                text = pytesseract.image_to_string(img)
                return text or ''
        except Exception as e:
            logger.warning(f"OCR failed on page {page_index+1}: {e}")
            return ''

    def _extract_tables_with_plumber(self, pdf_bytes: bytes, page_index: int) -> str:
        """Extract simple tables using pdfplumber if available."""
        if not PDFPLUMBER_AVAILABLE:
            return ''
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                if page_index < 0 or page_index >= len(pdf.pages):
                    return ''
                page = pdf.pages[page_index]
                tables = page.find_tables() if hasattr(page, 'find_tables') else page.extract_tables()
                rows_out = []
                for tbl in tables or []:
                    for row in tbl.rows if hasattr(tbl, 'rows') else tbl:
                        # row can be list[str] or cells; join with pipe delimiter
                        try:
                            cells = [c.get_text().strip() if hasattr(c, 'get_text') else (c or '').strip() for c in row]
                        except Exception:
                            cells = [str(c or '').strip() for c in row]
                        if any(cells):
                            rows_out.append(' | '.join(cells))
                return '\n'.join(rows_out)
        except Exception as e:
            logger.warning(f"Table extraction failed on page {page_index+1}: {e}")
            return ''

    def extract(self, file_content: bytes, filename: Optional[str] = None) -> ExtractionResult:
        """Extract text from PDF with page-level handling and graceful fallbacks."""
        if not PDF_AVAILABLE:
            return ExtractionResult(success=False, error="PyPDF2 not available for PDF extraction")

        start_time = datetime.now()
        warnings: List[str] = []

        try:
            # Check file size
            if len(file_content) > self.max_file_size:
                return ExtractionResult(
                    success=False,
                    error=f"File too large: {len(file_content)} bytes (max: {self.max_file_size})"
                )

            # Create PDF reader
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Handle encryption gracefully
            encrypted = False
            try:
                if getattr(pdf_reader, 'is_encrypted', False):
                    encrypted = True
                    try:
                        # Best-effort empty password decrypt
                        result = pdf_reader.decrypt("")
                        if result == 0:  # decrypt failed
                            return ExtractionResult(success=False, error="Encrypted PDF: decryption required")
                        else:
                            warnings.append("Encrypted PDF decrypted with empty password")
                    except Exception:
                        return ExtractionResult(success=False, error="Encrypted PDF: unable to decrypt")
            except Exception as e:
                warnings.append(f"Encryption check failed: {e}")

            # Extract text page by page with fallbacks
            page_count = len(pdf_reader.pages)
            text_content: List[str] = []
            ocr_used_pages: List[int] = []
            tables_extracted = 0

            for page_num, page in enumerate(pdf_reader.pages):
                page_text = ''
                try:
                    raw = page.extract_text() or ''
                    page_text = raw.strip()
                except Exception as e:
                    warnings.append(f"Failed PyPDF2 extract on page {page_num + 1}: {e}")
                    page_text = ''

                # If no text, try OCR fallback (if available)
                if not page_text:
                    ocr_text = self._ocr_page_with_plumber(file_content, page_num)
                    if ocr_text.strip():
                        page_text = ocr_text.strip()
                        ocr_used_pages.append(page_num + 1)
                    else:
                        warnings.append(f"No text found on page {page_num + 1}")

                # Attempt table extraction and append below text (if available)
                table_text = self._extract_tables_with_plumber(file_content, page_num)
                if table_text.strip():
                    page_text = (page_text + '\n' + table_text).strip() if page_text else table_text
                    tables_extracted += 1

                if page_text:
                    text_content.append(page_text)

            # Combine all text
            full_text = '\n\n'.join(text_content)
            cleaned_text = self._clean_text(full_text)

            if not cleaned_text:
                return ExtractionResult(success=False, error="No text content could be extracted from PDF")

            # Calculate extraction time
            extraction_time = (datetime.now() - start_time).total_seconds()

            # Get metadata
            metadata = {
                'filename': filename,
                'pages': page_count,
                'extraction_method': 'PyPDF2' + ('+pdfplumber/pytesseract' if ocr_used_pages or tables_extracted else ''),
                'encrypted': encrypted,
                'ocr_used_pages': ocr_used_pages,
                'tables_extracted': tables_extracted,
            }

            # Add PDF metadata if available
            try:
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                    })
            except Exception as e:
                warnings.append(f"Could not extract PDF metadata: {e}")

            stats = self._calculate_stats(cleaned_text)

            document = ExtractedDocument(
                content=cleaned_text,
                metadata=metadata,
                file_type='pdf',
                file_size=len(file_content),
                extraction_time=extraction_time,
                page_count=page_count,
                word_count=stats['word_count'],
                character_count=stats['character_count']
            )

            return ExtractionResult(success=True, document=document, warnings=warnings)

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ExtractionResult(success=False, error=f"PDF extraction failed: {str(e)}")

class DOCXExtractor(DocumentExtractor):
    """
    Extract text from DOCX files
    """

    def __init__(self):
        super().__init__()
        self.supported_types = ['docx']

    def extract(self, file_content: bytes, filename: Optional[str] = None) -> ExtractionResult:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            return ExtractionResult(
                success=False,
                error="python-docx not available for DOCX extraction"
            )

        start_time = datetime.now()
        warnings: List[str] = []

        try:
            # Check file size
            if len(file_content) > self.max_file_size:
                return ExtractionResult(
                    success=False,
                    error=f"File too large: {len(file_content)} bytes (max: {self.max_file_size})"
                )

            # Create document from bytes
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(' | '.join(row_text))

            # Combine all text
            all_text = paragraphs + table_text
            full_text = '\n\n'.join(all_text)
            cleaned_text = self._clean_text(full_text)

            if not cleaned_text:
                return ExtractionResult(
                    success=False,
                    error="No text content could be extracted from DOCX"
                )

            # Calculate extraction time
            extraction_time = (datetime.now() - start_time).total_seconds()

            # Get metadata
            metadata = {
                'filename': filename,
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables),
                'extraction_method': 'python-docx'
            }

            # Add document properties if available
            try:
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'category': core_props.category or '',
                    'comments': core_props.comments or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else ''
                })
            except Exception as e:
                warnings.append(f"Could not extract DOCX metadata: {e}")

            stats = self._calculate_stats(cleaned_text)

            document = ExtractedDocument(
                content=cleaned_text,
                metadata=metadata,
                file_type='docx',
                file_size=len(file_content),
                extraction_time=extraction_time,
                page_count=None,  # DOCX doesn't have fixed pages
                word_count=stats['word_count'],
                character_count=stats['character_count']
            )

            return ExtractionResult(
                success=True,
                document=document,
                warnings=warnings
            )

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"DOCX extraction failed: {str(e)}"
            )

class TextExtractor(DocumentExtractor):
    """
    Extract text from plain text files
    """

    def __init__(self):
        super().__init__()
        self.supported_types = ['txt', 'text', 'log', 'csv', 'json', 'xml', 'html']

    def extract(self, file_content: bytes, filename: Optional[str] = None) -> ExtractionResult:
        """Extract text from text files"""
        start_time = datetime.now()
        warnings: List[str] = []

        try:
            # Check file size
            if len(file_content) > self.max_file_size:
                return ExtractionResult(
                    success=False,
                    error=f"File too large: {len(file_content)} bytes (max: {self.max_file_size})"
                )

            # Detect encoding
            encoding = self._detect_encoding(file_content)

            try:
                text_content = file_content.decode(encoding)
            except UnicodeDecodeError:
                # Fallback to utf-8 with error handling
                text_content = file_content.decode('utf-8', errors='replace')
                warnings.append(f"Encoding detection failed, used UTF-8 with error replacement")

            # Clean text
            cleaned_text = self._clean_text(text_content)

            if not cleaned_text:
                return ExtractionResult(
                    success=False,
                    error="No text content found in file"
                )

            # Calculate extraction time
            extraction_time = (datetime.now() - start_time).total_seconds()

            # Determine file type from extension
            file_type = 'txt'
            if filename:
                ext = filename.split('.')[-1].lower()
                if ext in self.supported_types:
                    file_type = ext

            # Get metadata
            metadata = {
                'filename': filename,
                'encoding': encoding,
                'extraction_method': 'text_decode'
            }

            stats = self._calculate_stats(cleaned_text)

            document = ExtractedDocument(
                content=cleaned_text,
                metadata=metadata,
                file_type=file_type,
                file_size=len(file_content),
                extraction_time=extraction_time,
                page_count=None,
                word_count=stats['word_count'],
                character_count=stats['character_count']
            )

            return ExtractionResult(
                success=True,
                document=document,
                warnings=warnings
            )

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"Text extraction failed: {str(e)}"
            )

class MarkdownExtractor(DocumentExtractor):
    """
    Extract text from Markdown files
    """

    def __init__(self):
        super().__init__()
        self.supported_types = ['md', 'markdown']

    def extract(self, file_content: bytes, filename: Optional[str] = None) -> ExtractionResult:
        """Extract text from Markdown"""
        start_time = datetime.now()
        warnings: List[str] = []

        try:
            # Check file size
            if len(file_content) > self.max_file_size:
                return ExtractionResult(
                    success=False,
                    error=f"File too large: {len(file_content)} bytes (max: {self.max_file_size})"
                )

            # Decode text
            encoding = self._detect_encoding(file_content)
            try:
                markdown_content = file_content.decode(encoding)
            except UnicodeDecodeError:
                markdown_content = file_content.decode('utf-8', errors='replace')
                warnings.append("Encoding detection failed, used UTF-8 with error replacement")

            # Convert markdown to HTML and extract text
            if MARKDOWN_AVAILABLE:
                try:
                    html = markdown.markdown(markdown_content)
                    soup = BeautifulSoup(html, 'html.parser')
                    text_content = soup.get_text()
                except Exception as e:
                    warnings.append(f"Markdown parsing failed: {e}, using raw text")
                    text_content = markdown_content
            else:
                # Fallback: use raw markdown text
                text_content = markdown_content
                warnings.append("Markdown library not available, using raw text")

            # Clean text
            cleaned_text = self._clean_text(text_content)

            if not cleaned_text:
                return ExtractionResult(
                    success=False,
                    error="No text content found in Markdown file"
                )

            # Calculate extraction time
            extraction_time = (datetime.now() - start_time).total_seconds()

            # Get metadata
            metadata = {
                'filename': filename,
                'encoding': encoding,
                'extraction_method': 'markdown' if MARKDOWN_AVAILABLE else 'raw_text'
            }

            # Extract markdown headers for structure info
            headers = re.findall(r'^#+\s+(.+)$', markdown_content, re.MULTILINE)
            if headers:
                metadata['headers'] = headers[:10]  # First 10 headers
                metadata['header_count'] = len(headers)

            stats = self._calculate_stats(cleaned_text)

            document = ExtractedDocument(
                content=cleaned_text,
                metadata=metadata,
                file_type='markdown',
                file_size=len(file_content),
                extraction_time=extraction_time,
                page_count=None,
                word_count=stats['word_count'],
                character_count=stats['character_count']
            )

            return ExtractionResult(
                success=True,
                document=document,
                warnings=warnings
            )

        except Exception as e:
            logger.error(f"Markdown extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"Markdown extraction failed: {str(e)}"
            )

class DocumentProcessor:
    """
    Main document processor that coordinates all extractors
    """

    def __init__(self):
        self.extractors = {
            'pdf': PDFExtractor(),
            'docx': DOCXExtractor(),
            'txt': TextExtractor(),
            'markdown': MarkdownExtractor()
        }

        # File type mappings
        self.type_mappings = {
            'pdf': 'pdf',
            'docx': 'docx',
            'doc': 'docx',  # Treat DOC as DOCX (limited support)
            'txt': 'txt',
            'text': 'txt',
            'log': 'txt',
            'csv': 'txt',
            'json': 'txt',
            'xml': 'txt',
            'html': 'txt',
            'md': 'markdown',
            'markdown': 'markdown'
        }

    def get_file_type(self, filename: str, mime_type: Optional[str] = None) -> str:
        """Determine file type from filename and mime type"""
        if filename:
            ext = filename.split('.')[-1].lower()
            if ext in self.type_mappings:
                return self.type_mappings[ext]

        # Fallback to mime type
        if mime_type:
            if 'pdf' in mime_type:
                return 'pdf'
            elif 'word' in mime_type or 'officedocument' in mime_type:
                return 'docx'
            elif 'text' in mime_type:
                return 'txt'

        return 'txt'  # Default fallback

    def process_document(
        self,
        file_content: bytes,
        filename: Optional[str] = None,
        mime_type: Optional[str] = None
    ) -> ExtractionResult:
        """Process a document and extract its content"""
        try:
            # Determine file type
            file_type = self.get_file_type(filename, mime_type)

            # Get appropriate extractor
            extractor = self.extractors.get(file_type)
            if not extractor:
                return ExtractionResult(
                    success=False,
                    error=f"No extractor available for file type: {file_type}"
                )

            # Extract content
            result = extractor.extract(file_content, filename)

            # Add processing metadata
            if result.success and result.document:
                result.document.metadata['processor_version'] = '1.0'
                result.document.metadata['processed_at'] = datetime.now().isoformat()
                result.document.metadata['detected_type'] = file_type

            return result

        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"Document processing failed: {str(e)}"
            )

    def get_supported_types(self) -> List[str]:
        """Get list of supported file types"""
        return list(self.type_mappings.keys())

    def validate_file(self, file_content: bytes, filename: Optional[str] = None) -> Dict[str, Any]:
        """Validate file before processing"""
        validation = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'file_size': len(file_content),
            'detected_type': None
        }

        # Check file size
        max_size = 50 * 1024 * 1024  # 50MB
        if len(file_content) > max_size:
            validation['valid'] = False
            validation['errors'].append(f"File too large: {len(file_content)} bytes (max: {max_size})")

        # Check if file is empty
        if len(file_content) == 0:
            validation['valid'] = False
            validation['errors'].append("File is empty")

        # Detect file type
        if filename:
            file_type = self.get_file_type(filename)
            validation['detected_type'] = file_type

            if file_type not in self.extractors:
                validation['warnings'].append(f"File type '{file_type}' may not be fully supported")

        return validation

# Global instance
document_processor = DocumentProcessor()
