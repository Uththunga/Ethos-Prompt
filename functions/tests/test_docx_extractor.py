import io
import pytest
from docx import Document as DocxDocument

from src.rag.document_extractors import DOCX_AVAILABLE, DOCXExtractor

pytestmark = pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not available")


def make_docx(paragraphs=(), table_rows=None) -> bytes:
    d = DocxDocument()
    for p in paragraphs:
        d.add_paragraph(p)
    if table_rows:
        rows = len(table_rows)
        cols = len(table_rows[0])
        table = d.add_table(rows=rows, cols=cols)
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_docx_extractor_paragraphs_and_table():
    content = make_docx(paragraphs=["Hello", "World"], table_rows=[["A","B"],["C","D"]])
    res = DOCXExtractor().extract(content, filename="sample.docx")
    assert res.success is True
    assert res.document is not None
    meta = res.document.metadata
    assert meta.get("paragraphs") >= 2
    assert meta.get("tables") >= 1
    assert "Hello" in res.document.content
    assert "A | B" in res.document.content


def test_docx_extractor_no_text_failure():
    content = make_docx(paragraphs=[], table_rows=None)
    res = DOCXExtractor().extract(content, filename="empty.docx")
    assert res.success is False
    assert "No text content" in (res.error or "")

